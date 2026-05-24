"""Build VIDEO_SCRIPT.docx — the compressed 4-minute concept-video script.

Mirror of the speaker notes inside docs/RDTII_AutoExtract.pptx, packaged for use as a
teleprompter / for sharing with the team. Apache 2.0.
"""
from pathlib import Path
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = Path(__file__).parent
OUT = HERE / "VIDEO_SCRIPT.docx"
if len(sys.argv) > 1:
    OUT = HERE / sys.argv[1]

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)

# ---------- Title ----------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("RDTII AutoExtract — Concept Video Script")
tr.bold = True
tr.font.size = Pt(18)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("Team Arkova   ·   ~4 minutes   ·   ~556 spoken words at 140 wpm")
sr.italic = True
sr.font.size = Pt(11)

intro = doc.add_paragraph()
ir = intro.add_run(
    "Per hackathon guidelines, the video covers our strategy for automated discovery of "
    "legal documents and for mapping and verification of regulatory evidence."
)
ir.italic = True
ir.font.size = Pt(10)
ir.font.color.rgb = RGBColor(0x5F, 0x63, 0x68)

doc.add_paragraph()  # spacer


def beat(timing, slide, narration):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(2)
    r = h.add_run(f"{timing}   ·   Slide {slide}")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x0B, 0x3D, 0x91)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    nr = p.add_run(narration)
    nr.font.size = Pt(12)


beat("0:00 – 0:15", 1,
     "We are Team Arkova. We built RDTII AutoExtract — an open-source, "
     "model-agnostic pipeline that automates roughly eighty percent of the RDTII "
     "workflow for Pillars 6 and 7 across Asia-Pacific jurisdictions.")

beat("0:15 – 0:40", 2,
     "The problem we tackle is simple: today the RDTII workflow is almost "
     "entirely manual. ESCAP researchers search government portals, retrieve PDFs, "
     "read dense legal text, and extract structured fields — article by article, "
     "across many jurisdictions and languages. It does not scale, and existing "
     "literature targets the EU AI Act and GDPR, not RDTII.")

beat("0:40 – 1:00", 3,
     "Our six objectives follow directly. Automate discovery beyond the ESCAP "
     "database. Extract all six mandatory fields at article-level. Map every "
     "provision to Pillar 6 or 7 sub-indicators with verifiable citations. Ship a "
     "self-hostable open tool. Keep every AI component swappable. And provide a "
     "human-first review surface.")

beat("1:00 – 1:25", 4,
     "To make swappability real, we build on ports and adapters. The core domain "
     "depends only on interfaces — never on a concrete model. Every AI component "
     "— LLM, OCR, captioner, embeddings, vector store, graph library — is an "
     "adapter behind a port, changed via configuration. No vendor lock-in, no "
     "production proprietary dependency.")

beat("1:25 – 2:35", 5,
     "The pipeline runs in five stages. Stages one and two handle automated "
     "discovery of legal documents. A crawler — Playwright with anti-bot handling "
     "and archive fallback — locates regulations on national portals, gazettes, "
     "and ministry sites, recording full provenance per document. Then OCR "
     "converts every format to clean text at under five percent character error "
     "rate, and a vision-language captioner produces a short descriptor of each "
     "section. That caption becomes the basis for every downstream tag, so we are "
     "not hostage to OCR noise on scanned or non-English documents. "
     "Stage three handles mapping. We tag each section multi-label against "
     "indicator labels, then a RAG-based LLM call extracts the six mandatory "
     "fields with strict JSON schema and a source citation. Stage four is human "
     "verification — a non-technical reviewer console where any mapping is "
     "accepted, rejected, or edited in seconds, with a direct link back to the "
     "source span.")

beat("2:35 – 3:15", 6,
     "Stage five is our core contribution. From the tagged nodes, we derive two "
     "artifacts in parallel. The concept graph is subtractive — we start from the "
     "complete pairwise graph and drop edges below a calibrated threshold, "
     "leaving only topical borders. Then Louvain communities and weighted "
     "PageRank score importance within the surviving web. Independently, Formal "
     "Concept Analysis on the node-by-tag matrix produces a generality tree: "
     "more tags imply fewer matching sections, so depth equals specificity. "
     "Together they give us entry-point categories opening into ranked "
     "subcategories — navigable cross-jurisdiction evidence.")

beat("3:15 – 3:35", 7,
     "Every model named here is a starting suggestion, not a commitment. "
     "Development uses Claude or GPT; production resolves to open-weight "
     "Llama 3.1 via Ollama with a single config change. All reused components "
     "are Apache-2.0 compatible, audited for license discipline.")

beat("3:35 – 3:50", 8,
     "On viability — open-weight self-hosted cost is roughly five to ten cents "
     "per fifty-page document, API at ten to twenty-five cents. We are "
     "finale-ready for ten countries with no retraining. Everything ships via "
     "docker-compose on commodity hardware.")

beat("3:50 – 3:55", 9,
     "The path to October runs from the May application through training "
     "workshops, Round 1, the hybrid pitch, and the Bangkok finale.")

beat("3:55 – 4:00", 10,
     "To close: not a new model — a complete, deployable, open-source RDTII "
     "pipeline that does not yet exist. Auditable, multilingual, and built for "
     "the people who do the work. Thank you.")

doc.save(OUT)
print(f"wrote {OUT}")
