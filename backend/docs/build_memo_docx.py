"""Build TECHNICAL_MEMO.docx from TECHNICAL_MEMO.md (2-page submission)."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = Path(__file__).parent
OUT = HERE / "TECHNICAL_MEMO.docx"

doc = Document()

# Tight page setup so it fits 2 pages
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)


def add_heading(text, size=14, bold=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p


def add_para(runs, size=10, italic=False, align=None):
    """runs = list of (text, bold) tuples, or a plain string."""
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if isinstance(runs, str):
        runs = [(runs, False)]
    for text, bold in runs:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
    return p


def add_mono(text, size=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(size)
    return p


def add_bullet(runs, size=10):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if isinstance(runs, str):
        runs = [(runs, False)]
    for text, bold in runs:
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
    return p


# ---------- Title ----------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
tr = title.add_run("Technical Memo — RDTII AutoExtract")
tr.bold = True
tr.font.size = Pt(16)

sub = doc.add_paragraph()
sr = sub.add_run(
    "Global Hackathon on AI for Digital Trade Regulatory Analysis "
    "(UN ESCAP & KMITL, 2026). Apache 2.0."
)
sr.italic = True
sr.font.size = Pt(9)

sub2 = doc.add_paragraph()
sr2 = sub2.add_run("Companion deck covers walkthrough; this memo = 2-page technical summary.")
sr2.italic = True
sr2.font.size = Pt(9)

# ---------- Problem ----------
add_para([
    ("Problem. ", True),
    ("Automate ~80% of the RDTII workflow (discover → describe) for ", False),
    ("Pillar 6 (Cross-border Data Flows) ", True),
    ("and ", False),
    ("Pillar 7 (Domestic Data Protection) ", True),
    ("at article-level granularity, with a transparent 20% human-review step. ", False),
    ("Model-agnostic by design", True),
    (" — every model is a swappable adapter.", False),
])

# ---------- Pipeline ----------
add_heading("Pipeline", size=12)
pipeline = (
    "┌──────────┐   ┌───────────────────────┐   ┌──────────────┐   ┌──────────────┐   ┌───────────────────┐\n"
    "│ 1 DISCOVER│──▶│ 2 OCR + SigLIP CAPTION │──▶│ 3 TAG+EXTRACT │──▶│ 4 HUMAN REVIEW│──▶│ 5 CONCEPT GRAPH    │\n"
    "│ crawl,    │   │ <5% CER; caption =     │   │ multi-label   │   │ accept/reject │   │ weighted edges →   │\n"
    "│ anti-bot, │   │ BASIS for tags; VLM    │   │ + 6 fields +  │   │ /edit; audit  │   │ prune(θ) →         │\n"
    "│ provenance│   │ on scanned/non-EN docs │   │ indicator map │   │ view, export  │   │ community → FCA+PR │\n"
    "└──────────┘   └───────────────────────┘   └──────────────┘   └──────────────┘   └───────────────────┘\n"
    "   DocumentSource      OCREngine/Captioner       Tagger/LLM         FindingRepo      GraphBuilder/Ranker"
)
add_mono(pipeline, size=7)

add_para([
    ("6 mandatory fields / article: ", True),
    ("title · last_update · url · scope · provisions · impact ", False),
    ("(+ pillar, indicator, confidence, review_status). Document-level summaries rejected.", False),
])

add_para([
    ("Stage 5 (core contribution): ", True),
    ("tagged sections = nodes; edges = IDF-weighted tag overlap + embedding cosine; pruned at calibrated ", False),
    ("θ ", True),
    ("(fixed seeds + θ ⇒ reproducible). Community detection + ", False),
    ("Formal Concept Analysis ", True),
    ("lattice + weighted ", False),
    ("PageRank ", True),
    ("→ navigable entry-point → sub-topic hierarchy for cross-jurisdiction evidence discovery.", False),
])

# ---------- Architecture ----------
add_heading("Architecture — ports & adapters", size=12)
add_para("Core domain imports no framework/model SDK; concrete tools are config-selected adapters.")

rows = [
    ("Port", "Suggested adapter (dev)", "Open-weight target", "License"),
    ("DocumentSource", "Playwright + BeautifulSoup", "same", "BSD/MIT"),
    ("OCREngine", "Tesseract 5 + OpenCV / PaddleOCR", "same", "Apache-2.0"),
    ("Captioner", "Qwen2-VL / ColQwen2", "same", "Apache-2.0"),
    ("Tagger (embeddings)", "BGE-M3 + reranker", "same", "MIT/Apache"),
    ("VectorStore", "pgvector / Chroma", "same", "PostgreSQL/Apache"),
    ("LLMProvider", "Claude / GPT", "Llama 3.1 8B/70B via Ollama", "Apache-compat"),
    ("GraphBuilder/Ranker", "NetworkX + FCA `concepts`", "same", "BSD/MIT"),
]
table = doc.add_table(rows=len(rows), cols=4)
table.style = "Light Grid Accent 1"
for i, row in enumerate(rows):
    for j, text in enumerate(row):
        cell = table.cell(i, j)
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.font.size = Pt(8.5)
        if i == 0:
            r.bold = True

warn = doc.add_paragraph()
warn.paragraph_format.space_before = Pt(4)
wr = warn.add_run(
    "⚠ License discipline: ColQwen2 over PaliGemma/ColPali (Gemma); NetworkX/Louvain over GPL "
    "leidenalg/igraph; no CC-BY-NC models. All components Apache-2.0-compatible."
)
wr.italic = True
wr.font.size = Pt(9)

# ---------- Cost ----------
add_heading("Cost per 50-page document (preliminary)", size=12)
add_para("~200 article-level chunks / 50-pp doc.", size=10)

cost_rows = [
    ("Configuration", "Total / 50 pp"),
    ("Open-weight (Llama 3.1 8B, self-host GPU)", "~USD 0.05–0.10 (API: $0.00)"),
    ("API (Claude/GPT)", "~USD 0.10–0.25"),
]
ct = doc.add_table(rows=len(cost_rows), cols=2)
ct.style = "Light Grid Accent 1"
for i, row in enumerate(cost_rows):
    for j, text in enumerate(row):
        cell = ct.cell(i, j)
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.font.size = Pt(9)
        if i == 0:
            r.bold = True

add_para(
    "Open-weight basis: A10G-class GPU at batch rates. Final CER / latency / cost reported post-test.",
    size=9,
)

# ---------- Generalisation ----------
add_heading("Generalisation, fine-tuning, originality", size=12)
add_bullet([
    ("Generalisation: ", True),
    ("no per-country retraining; non-English via captioning + translation; scanned PDFs via OCR "
     "(<5% CER); anti-bot portals via headless browser + archive fallback. Target ≥3 of 10 "
     "assigned countries at finale.", False),
])
add_bullet([
    ("Fine-tuning (planned): ", True),
    ("few-shot tune small encoder/classifier for Pillar 6/7 sub-indicators using RDTII taxonomy; "
     "θ calibrated F1-optimal on RDTII labelled data. Weights published in repo.", False),
])
add_bullet([
    ("Originality: ", True),
    ("not a new model — a complete, deployable, open-source RDTII pipeline that does not yet "
     "exist, with a concept-graph layer for cross-jurisdiction evidence discovery.", False),
])

# ---------- Deploy ----------
add_para([
    ("Deploy: ", True),
    ("docker-compose up → api + web + Postgres/pgvector. No proprietary production dependency. "
     "Full design: docs/ARCHITECTURE.md, docs/GRAPH_PIPELINE.md.", False),
])

doc.save(OUT)
print(f"wrote {OUT}")
