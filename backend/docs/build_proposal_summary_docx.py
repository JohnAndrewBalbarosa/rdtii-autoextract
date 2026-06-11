"""Build PROPOSAL_SUMMARY.docx — Title + 200-word summary + 200-word problem/objectives."""
from pathlib import Path
import sys
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = Path(__file__).parent
OUT = HERE / "PROPOSAL_SUMMARY.docx"
if len(sys.argv) > 1:
    OUT = HERE / sys.argv[1]

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)


def heading(text, size=13):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)


def body(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.size = Pt(11)


# ---------- Title block ----------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("Zetarix")
tr.bold = True
tr.font.size = Pt(18)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run(
    "An Open-Source AI Pipeline for Automated Digital Trade Regulatory Analysis "
    "Across Asia-Pacific Jurisdictions"
)
sr.italic = True
sr.font.size = Pt(12)

attr = doc.add_paragraph()
attr.alignment = WD_ALIGN_PARAGRAPH.CENTER
ar = attr.add_run(
    "Team Arkova · Global Hackathon on AI for Digital Trade Regulatory Analysis "
    "(UN ESCAP & KMITL, 2026) · Apache 2.0"
)
ar.font.size = Pt(10)
ar.italic = True

# ---------- Proposal Summary ----------
heading("Proposal Summary")
body(
    "Zetarix is an open-source, model-agnostic AI pipeline that automates roughly "
    "80% of the UN ESCAP Regional Digital Trade Integration Index workflow — discovery "
    "through description — for Pillar 6 (Cross-Border Data Flows) and Pillar 7 (Domestic "
    "Data Protection) across Asia-Pacific jurisdictions. The system runs five composable "
    "stages: crawler-based document discovery with compliant access handling and provenance "
    "tracking; OCR with vision-language captioning for scanned, non-English, and visually "
    "laid-out texts, targeting under five percent character error rate; LLM-driven "
    "multi-label tagging and article-level extraction of the six mandatory RDTII fields "
    "with verifiable source citations; a non-technical human-review console where any "
    "mapping can be verified, rejected, or edited in seconds; and a concept-graph layer "
    "that links related provisions across jurisdictions via threshold-pruned weighted "
    "edges, Louvain community detection, Formal Concept Analysis lattices, and weighted "
    "PageRank for navigable cross-jurisdiction evidence discovery. Built on a "
    "ports-and-adapters architecture, every model — LLM, OCR engine, captioner, embedder, "
    "vector store, and graph library — is a swappable adapter changed via configuration "
    "without touching the domain. The reference stack uses Apache-2.0-compatible "
    "components only; production targets open-weight Llama 3.1 via Ollama at an estimated "
    "cost of USD 0.05 to 0.10 per fifty-page document. The full system is self-hostable "
    "through docker-compose with no proprietary production dependency."
)

# ---------- Problem Understanding and Objectives ----------
heading("Problem Understanding and Objectives")
body(
    "The current RDTII workflow requires ESCAP researchers to manually search government "
    "portals, retrieve regulatory documents, read dense legal text, and extract structured "
    "data at article-level granularity across multiple Asia-Pacific jurisdictions. The "
    "process is time-intensive, difficult to scale, and bottlenecked by researcher "
    "availability, messy portals, and language coverage spanning English, Vietnamese, "
    "Thai, Bahasa Malaysia, and others. Existing automated approaches in the literature "
    "concentrate almost entirely on EU instruments such as GDPR and the AI Act; no "
    "open-source pipeline today targets the RDTII indicators end-to-end or handles "
    "Asia-Pacific languages, scanned PDFs, and document conventions in one deployable "
    "system. Our objectives are: automate discovery, retrieval, OCR, captioning, tagging, "
    "and extraction so researchers receive review-ready findings rather than raw "
    "documents; deliver article-level outputs containing the six mandatory fields — "
    "title, last update, URL, scope, provisions, impact — with verifiable mapping to "
    "Pillar 6 and 7 sub-indicators and per-record confidence scores; preserve a "
    "transparent human-review surface where any mapping is verified, rejected, or edited "
    "in seconds by non-technical users; maintain strict model-agnosticism through ports "
    "and adapters so open-weight upgrades require only configuration changes; demonstrate "
    "generalisation across at least three of ten assigned countries under live "
    "conditions; and ship a self-hostable Apache-2.0 reference implementation that any "
    "policy team can adopt, audit, and extend without proprietary infrastructure."
)

doc.save(OUT)
print(f"wrote {OUT}")
